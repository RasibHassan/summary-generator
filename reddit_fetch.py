import praw
import re
from docx import Document
import os
# Reddit API setup
reddit = praw.Reddit(
    client_id="neOh9sh9KTbFdYAD8EhXyA",
    client_secret="rZePjcmM9BjI5ZtGN7LM-mL4_sOKqg",
    user_agent="script:reddit_fetcher:1.0 (by u/Famous_Ratio_7003)",
    username="Famous_Ratio_7003",
    password="hasan231"
)

# Your list of Reddit post URLs

def extract_post_id(url):
    """Extract post ID from Reddit URL"""
    match = re.search(r'/comments/([a-z0-9]+)/', url)
    return match.group(1) if match else None

def main(urls):
    """Main function to process Reddit URLs"""
    print("📥 Fetching Reddit posts and comments...")

    doc = Document()
    doc.add_heading("Reddit PMP Posts and Comments", 0)

    for url in urls:
        post_id = extract_post_id(url)
        if post_id:
            submission = reddit.submission(id=post_id)
            submission.comments.replace_more(limit=None)
            
            # Add post title and body
            doc.add_heading(submission.title, level=1)
            doc.add_paragraph(submission.selftext)

            # Add comments
            doc.add_heading("Comments", level=2)
            for comment in submission.comments.list():
                doc.add_paragraph(f"[Comment] {comment.body}", style='Intense Quote')

    output_path = os.path.join("data", "reddit_combined_text.docx")
    doc.save(output_path)
    print(f"✅ Reddit post content and comments saved to '{output_path}'")