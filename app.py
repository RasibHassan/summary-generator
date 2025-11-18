import os
import streamlit as st
import pandas as pd
import json
from docx import Document
from youtube_summarizer import summarize_youtube_video
from document_summarizer import generate_summary_from_file
from video_summarizer import get_summary_from_video
from datetime import datetime
from openai import OpenAI
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from tavily import TavilyClient
from youtube_fetch import main as youtube_main
from reddit_fetch import main as reddit_main
from urls_fetch import main as urls_main
from gpt import main as gpt_main
from claude import main as claude_main
from dotenv import load_dotenv

load_dotenv()
# Configuration
os.environ["STREAMLIT_WATCH_USE_POLLING"] = "true"

# API Keys and Clients
openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key)
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
tavily_client = TavilyClient(api_key=TAVILY_API_KEY)

# Setup directories
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
uploaded_filenames = []

# Helper Functions for Summary Generator
def add_formatted_run(paragraph, text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*\*)')
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if pos < start:
            paragraph.add_run(text[pos:start])
        matched_text = match.group()
        cleaned_text = matched_text.strip('*')
        run = paragraph.add_run(cleaned_text)
        run.bold = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])

def write_to_word(text, output_file="generated_summary.docx"):
    doc = Document()
    title = doc.add_heading("Generated Topic-Wise Plan", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^Chapter \d+: ", line, re.IGNORECASE):
            doc.add_heading(line, level=2)
        elif line.startswith(("-", "*", "•")):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, line[1:].strip())
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, line.strip())
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, line)

    doc.save(output_file)

# # Helper Functions for Study Plan Generator------with Filtering links on cosine similarity
# def search_and_extract(prompt, category, include_domains=None, exclude_domains=None):
#     max_result = 10  # Limit to 10 results
#     if category == "YouTube Videos":
#         max_result = 5

#     response = tavily_client.search(
#         query=prompt,
#         search_depth="advanced",
#         max_results=max_result,
#         include_answer='advanced',
#         # include_raw_content=False,------by Owais
#         include_raw_content=True,
#         include_images=False,
#         include_domains=include_domains,
#         exclude_domains=exclude_domains
#     )
#     return [{"category": category, "url": r["url"]} for r in response.get("results", [])],response.get("answer", "")

from tavily import TavilyClient
from sentence_transformers import SentenceTransformer, util
import numpy as np
from pytube import YouTube

def filter_youtube_by_length(urls, max_minutes=60):
    """
    Filters out YouTube URLs whose duration is longer than max_minutes.
    Returns a list of allowed URLs.
    """
    allowed_urls = []
    for url in urls:
        try:
            yt = YouTube(url)
            length_minutes = yt.length / 60  # yt.length is in seconds
            if length_minutes <= max_minutes:
                allowed_urls.append(url)
            else:
                st.warning(f"⚠ Skipped '{yt.title}' ({length_minutes:.1f} min) — longer than {max_minutes} minutes.")
        except Exception as e:
            st.warning(f"⚠ Could not fetch video info for {url}: {e}")
    return allowed_urls
# Load embedding model once
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

def embed_text(text):
    """Convert text to embedding."""
    return embedding_model.encode(text, convert_to_tensor=True)

def cosine_similarity(a, b):
    """Compute cosine similarity between two embeddings."""
    return util.cos_sim(a, b).item()

def search_and_extract(prompt, category, topic, similarity_threshold=0.60,
                       include_domains=None, exclude_domains=None, max_results=6):

    # Step 1: Tavily API Call
    response = tavily_client.search(
        query=prompt,
        search_depth="advanced",
        max_results=max_results,
        include_answer='advanced',
        include_raw_content=True,      # <-- Key for accurate similarity
        include_images=False,
        include_domains=include_domains,
        exclude_domains=exclude_domains
    )

    results = response.get("results", [])
    summary_answer = response.get("answer", "")

    if not results:
        return [], summary_answer

    # Step 2: Embed the main topic
    candidate_texts = []    
    for r in results:
        # Combine all available metadata text
        combined = "".join([
            r.get("title") or " ",
            r.get("snippet") or " ",
            r.get("raw_content") or " ",
            r.get("url") or ""
        ])
        candidate_texts.append(combined)
    url_embeddings = embedding_model.encode(candidate_texts, convert_to_tensor=True, batch_size=8)

    # step 4 topic encoding
    topic_embedding = embedding_model.encode(topic, convert_to_tensor=True)
    # step 5 compute batch cosine similarity
    similarities = util.cos_sim(topic_embedding, url_embeddings)[0]
    # Step 6: Filter based on similarity threshold
    filtered_results = []
    for i, sim in enumerate(similarities):
        if sim >= similarity_threshold:
            filtered_results.append({
                "category": category,
                "url": results[i].get("url"),
                "title": results[i].get("title", ""),
                "similarity": round(float(sim), 3)
            })

    return filtered_results, summary_answer


# App Configuration
st.set_page_config(page_title="Smart Academic Assistant", layout="centered")

# Main App Title and Feature Selection
st.title("🎓 Smart Academic Assistant")
st.markdown("Choose between generating summaries from videos/documents or creating study plans for certifications.")

# Feature Toggle
feature_choice = st.radio(
    "Select Feature:",
    ["📝 Summary Generator", "🎯 Study Plan Generator"],
    horizontal=True
)

st.divider()

# FEATURE 1: SUMMARY GENERATOR
if feature_choice == "📝 Summary Generator":
    st.title("🎥 Smart Summary Generator")
    st.markdown("Start by uploading video files to generate summaries. Optionally, you can also add YouTube URLs and documents.")

    # Initialize session state for summary generator
    if "summary_generated" not in st.session_state:
        st.session_state.summary_generated = False
    if "summary_path" not in st.session_state:
        st.session_state.summary_path = ""
    if "timestamp" not in st.session_state:
        st.session_state.timestamp = ""
    
    # Clear study plan session states when in summary mode
    if "search_results" in st.session_state:
        del st.session_state["search_results"]
    if "answers" in st.session_state:
        del st.session_state["answers"]

    # === Primary Input: Video Upload ===
    st.subheader("🎥 Upload Video Files")
    video_files = st.file_uploader("Upload videos", type=["mp4", "mov", "avi"], accept_multiple_files=True)

    # === Optional Input: YouTube URLs ===
    use_youtube = st.checkbox("📺 I want to add YouTube Video URLs")
    youtube_urls = []
    if use_youtube:
        st.subheader("📺 YouTube Video URLs")
        youtube_input = st.text_area("Enter one YouTube URL per line", placeholder="https://www.youtube.com/watch?v=...")
        youtube_urls = youtube_input.strip().splitlines()

    # === Optional Input: Document Upload ===
    use_docs = st.checkbox("📄 I want to upload documents (PDF/DOCX)")
    document_files = []
    if use_docs:
        st.subheader("📄 Upload Document Files")
        document_files = st.file_uploader("Upload documents", type=["pdf", "docx"], accept_multiple_files=True)
    if youtube_urls:
        st.subheader("📺 YouTube Video URLs")
        youtube_urls = youtube_input.strip().splitlines()
        
        # Filter out videos longer than 1 hour
        youtube_urls = filter_youtube_by_length(youtube_urls, max_minutes=40)

    # === Generate Summaries ===
    if st.button("🚀 Generate Summary"):
        doc = Document()
        doc.add_heading("Merged Summaries", level=1)
        added_any_summary = False

        for video_file in video_files:
            video_path = os.path.join(UPLOAD_DIR, video_file.name)
            with open(video_path, "wb") as f:
                f.write(video_file.read())
            st.write(f"🔄 Generating summary for video: {video_file.name}")
            summary = get_summary_from_video(video_path)
            if summary:
                doc.add_heading("Video File Summary", level=2)
                doc.add_heading(f"File: {video_file.name}", level=2)
                doc.add_paragraph(summary)
                added_any_summary = True
                st.success(f"✔ Summary added for video: {video_file.name}")
            else:
                st.warning(f"⚠ No summary generated for video: {video_file.name}")

        count = 1
        for url in youtube_urls:
            if url.strip():
                st.write(f"🔄 Generating summary for YouTube URL-{count}")
                summary = summarize_youtube_video(url)
                if summary:
                    doc.add_heading("YouTube Video Summary", level=2)
                    doc.add_heading(f"URL: {url}", level=2)
                    doc.add_paragraph(summary)
                    added_any_summary = True
                    st.success(f"✔ Summary added for: URL-{count}")
                    count += 1
                else:
                    st.warning(f"⚠ No summary for: {url}")

        for doc_file in document_files:
            doc_path = os.path.join(UPLOAD_DIR, doc_file.name)
            with open(doc_path, "wb") as f:
                f.write(doc_file.read())
            st.write(f"🔄 Generating summary for document: {doc_file.name}")
            summary = generate_summary_from_file(doc_path)
            if summary:
                doc.add_heading("Document Summary", level=2)
                doc.add_heading(f"File: {doc_file.name}", level=2)
                doc.add_paragraph(summary)
                added_any_summary = True
                st.success(f"✔ Summary added for document: {doc_file.name}")
            else:
                st.warning(f"⚠ No summary generated for document: {doc_file.name}")

        if added_any_summary:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join(UPLOAD_DIR, f"merged_summary_{timestamp}.docx")
            doc.save(output_path)

            # Store in session state
            st.session_state.summary_generated = True
            st.session_state.summary_path = output_path
            st.session_state.timestamp = timestamp
            st.success("✅ Summary generation complete!")

    # === GPT Reformatting ===
    if st.session_state.summary_generated:
        st.subheader("✨ Enhance Summary with GPT")
        selected_model = st.selectbox(
                "Choose GPT Model",
                options=["gpt-4.1-mini", "gpt-5-mini", "gpt-5"],
                index=0
            )


        if st.button("🔄 Reformat Using GPT"):
            with st.spinner("🧠 GPT is processing and enhancing the summary..."):
                try:
                    with open(st.session_state.summary_path, "rb") as f:
                        docx_file = Document(f)
                        full_text = "\n".join([para.text for para in docx_file.paragraphs if para.text.strip()])

                    if not full_text.strip():
                        st.error("❌ Summary content is empty. Cannot process.")
                    else:
                        prompt = f"""
        You are an expert tutor. Merge these study materials summaries into one complete learning guide.

        TASK:
        1. Combine all notes/chapters/topics into a single, organized summary
        2. Arrange topics in logical learning order (basics → advanced)
        3. Include ALL points from every source
        4. Highlight any conflicts between sources and explain why they exist
        5. Add your own knowledge to fill gaps and provide context
        6. Ensure the final summary is comprehensive and easy to follow

        note: generate a detailed summary as much as tokens allow 

        STRUCTURE:
        ## FUNDAMENTALS
        Core concepts and definitions

        ## KEY TOPICS  
        Main ideas with examples and applications

        ## ADVANCED CONCEPTS
        Complex topics and expert insights

        ## SUMMARY
        Key takeaways and important points to remember

        Make it easy to understand and study from. Don't skip anything.

        Study Material summaries:
        {full_text}
        Now write the final, detailed, well-organized explanation """

                        if selected_model == "gpt-4.1-mini":
                            response = client.chat.completions.create(
                                model="gpt-4.1-mini",
                                messages=[
                                    {"role": "system", "content": "You are an expert summarizer and tutor."},
                                    {"role": "user", "content": prompt}
                                ],
                                max_tokens=10000,
                                temperature=0.5
                            )
                            formatted_content = response.choices[0].message.content.strip()
                        else:
                            response = client.responses.create(
                                        model=selected_model,
                                        input=[{"role": "user", "content": prompt}],
                                        text={"verbosity": "high"},
                                        reasoning={"effort": "low"},

                                    )

                            formatted_content = response.output_text
                        gpt_path = os.path.join(UPLOAD_DIR, f"enhanced_summary.docx")

                        write_to_word(formatted_content, gpt_path)

                        with open(gpt_path, "rb") as f:
                            st.download_button(
                                label="📥 Download GPT-Enhanced Summary",
                                data=f,
                                file_name="Formatted_summary.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        st.success("✅ GPT-enhanced summary is ready!")

                        try:
                            for filename in os.listdir(UPLOAD_DIR):
                                file_path = os.path.join(UPLOAD_DIR, filename)
                                if os.path.isfile(file_path):
                                    os.remove(file_path)
                            st.info("🧹 Temporary uploaded files have been deleted.")
                        except Exception as e:
                            st.warning(f"⚠ Error during cleanup: {e}")
                except Exception as e:
                    st.error(f"❌ GPT API Error: {e}")

# FEATURE 2: STUDY PLAN GENERATOR
elif feature_choice == "🎯 Study Plan Generator":
    st.title("🎓 Certification Resource Finder")

    # Initialize session state for study plan generator
    if "uploaded_files" not in st.session_state:
        st.session_state.uploaded_files = []
    
    # Clear summary session states when in study plan mode
    if "summary_generated" in st.session_state:
        st.session_state.summary_generated = False
    if "summary_path" in st.session_state:
        st.session_state.summary_path = ""
    if "timestamp" in st.session_state:
        st.session_state.timestamp = ""

    # Input field for certificate/topic
    user_input = st.text_input("Enter the certification or topic name", placeholder="e.g., PMI-PBA")

    st.subheader("🎨 Customize Your Study Plan")

    theme = st.text_input(
    "Select a theme for output:",placeholder="e.g., A clean dark slide theme with bold typography.")

    font = st.selectbox(
    "Select font style:",
    ["Sans-serif", "Serif", "Mono", "Open Sans", "Roboto", "Inter", "Georgia"]
    )

    detail_level = st.selectbox(
    "Choose level of detail:",
    ["Basic", "Medium", "Advanced", "Very Detailed"]
    )

    study_duration = st.number_input(
    "How many weeks should the study plan cover?",
    min_value=1, max_value=52, value=4
    )

    tone = st.selectbox(
    "Select writing tone:",
    ["Professional", "Friendly", "Motivational", "Minimalistic", "Teaching Style", "Exam-focused"]
    )


    # Tavily search
    if st.button("Search Resources") and user_input:
        with st.spinner("Searching..."):
            
            all_results = []

            yt_prompt = f'Find high-quality {user_input} preparation guide and tutorial videos. Focus on complete guides, exam tips, and step-by-step walkthroughs from popular educational channels. Do not include playlists or channel links'
            # yt_urls, yt_answer = search_and_extract(yt_prompt, "YouTube Videos", include_domains=["youtube.com"])
            yt_urls, yt_answer = search_and_extract(yt_prompt, "YouTube Videos",topic=user_input,similarity_threshold=0.50, include_domains=["youtube.com"])
            all_results += yt_urls

            reddit_prompt = f'Find the Reddit posts discussing preparation strategies, shared experiences, focus areas, and recommended resources for the {user_input}.'
            # reddit_urls, reddit_answer = search_and_extract(reddit_prompt, "Reddit Posts", include_domains=["reddit.com"])
            reddit_urls, reddit_answer = search_and_extract(reddit_prompt, "Reddit Posts",topic=user_input,similarity_threshold=0.50, include_domains=["reddit.com"],max_results=20)
            all_results += reddit_urls

            web_prompt = f'Find the best {user_input} study guides, preparation articles, and detailed learning resources from 2024. Include study plans, topic breakdowns, expert tips, and recommended materials.'
            
            # web_urls, web_answer = search_and_extract(web_prompt, "Web Resources", exclude_domains=["youtube.com", "reddit.com"])
            web_urls, web_answer = search_and_extract(web_prompt, "Web Resources",topic=user_input,similarity_threshold=0.50, exclude_domains=["youtube.com", "reddit.com"])
            all_results += web_urls

            st.session_state["answers"] = {
                "web": web_answer
            }

            df = pd.DataFrame(all_results)
            st.session_state["search_results"] = df

    # Display results
    if "search_results" in st.session_state:
        st.subheader("🔗 Tavily Search Results")
        df = st.session_state["search_results"]
        for category in df["category"].unique():
            st.markdown(f"**{category}**")
            for _, row in df[df["category"] == category].iterrows():
                st.markdown(f"- [{row['url']}]({row['url']})")

        st.divider()
        st.subheader("➕ Add More Links Manually (comma-separated)")

        # Text areas for additional links
        more_youtube = st.text_area("More YouTube Links", placeholder="https://youtube.com/xyz1, https://youtube.com/xyz2")
        more_reddit = st.text_area("More Reddit Links", placeholder="https://reddit.com/abc1, https://reddit.com/abc2")
        more_other = st.text_area("More Other Links", placeholder="https://example.com/article1, https://blog.com/post2")

        # File Upload
        st.divider()
        st.subheader("📤 Upload Supporting Files")
        files = st.file_uploader("Upload PDF or DOCX files", type=["pdf", "docx"], accept_multiple_files=True)

        if files:
            os.makedirs("data", exist_ok=True)

            for f in files:
                file_path = os.path.join("data", f.name)
                with open(file_path, "wb") as out_file:
                    out_file.write(f.read())
                uploaded_filenames.append(f.name)
                st.markdown(f"- ✅ Saved: `{file_path}`")

            st.session_state.uploaded_files = uploaded_filenames

        # Submit and Save
        st.divider()
        if st.button("✅ Submit "):
            # Combine Tavily + manual
            def parse_links(text): return [url.strip() for url in text.split(",") if url.strip()]
            final_data = {
                "youtube": [*df[df["category"] == "YouTube Videos"]["url"].tolist()] + parse_links(more_youtube),
                "reddit": [*df[df["category"] == "Reddit Posts"]["url"].tolist()] + parse_links(more_reddit),
                "other": [*df[df["category"] == "Web Resources"]["url"].tolist()] + parse_links(more_other),
                "files": st.session_state.uploaded_files,
                "answers": st.session_state["answers"]  # Add this line
            }

            # ---- Run the fetchers after saving ----
            st.divider()
            st.subheader("🚀 Running Data Fetch Modules")

            # Ensure folders exist
            os.makedirs("logs", exist_ok=True)
            os.makedirs("data", exist_ok=True)

            # Call YouTube Fetch
            st.write("▶️ Running YouTube transcript fetcher...")
            youtube_main(final_data["youtube"])

            # Call Reddit Fetch
            st.write("▶️ Running Reddit post fetcher...")
            reddit_main(final_data["reddit"])

            # Call Tavily Web Extractor
            st.write("▶️ Running Webpage text extractor...")
            urls_main(final_data["other"],final_data["answers"]["web"])

            st.write("🧠 Generating topic insights with GPT...")
            gpt_main(certificate_name=user_input)
            
            # === Run CLAUDE.py ===
            st.write("📚 Generating final study plan with Claude...")
            # claude_main()------by ow

            claude_main(
               topic=user_input,
               theme=theme,
               font=font,
               detail_level=detail_level,
              duration=study_duration,
              tone=tone
            )


            # Final Output Info
            st.success("📝 Study Plan Generated: `STUDY_GUIDE.HTML.docx`")

        # Placeholder for future output
        st.subheader("📦 Final Output")

        html_path = "study_guide.html"
        if os.path.exists(html_path):
            with open(html_path, "r", encoding="utf-8") as html_file:
                st.download_button(
                    label="📥 Download Study Plan (HTML)",
                    data=html_file.read(),
                    file_name="study_guide.html",
                    mime="text/html"
                )
