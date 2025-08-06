from youtube_transcript_api import YouTubeTranscriptApi, NoTranscriptFound, VideoUnavailable
from urllib.parse import urlparse, parse_qs
from docx import Document
import time
import random

def extract_video_id(url):
    """Extract video ID from YouTube URL"""
    parsed_url = urlparse(url)
    if 'youtu.be' in parsed_url.netloc:
        return parsed_url.path.lstrip('/')
    elif 'youtube.com' in parsed_url.netloc:
        return parse_qs(parsed_url.query).get('v', [None])[0]
    return None

def get_transcript_with_rate_limit(video_id, delay=3, max_retries=3):
    """Extract transcript with rate limiting and retry logic"""
    
    for attempt in range(max_retries):
        try:
            print(f"    Attempt {attempt + 1}/{max_retries}")
            
            # Add random delay to make requests less predictable
            actual_delay = delay + random.uniform(0, 2)
            print(f"    Waiting {actual_delay:.1f} seconds...")
            time.sleep(actual_delay)
            
            # Try to get transcript
            transcript = YouTubeTranscriptApi.get_transcript(video_id)
            transcript_text = ' '.join([entry['text'] for entry in transcript])
            
            print(f"    ✅ Success! Got {len(transcript_text)} characters")
            return transcript_text
            
        except (NoTranscriptFound, VideoUnavailable) as e:
            print(f"    ❌ No transcript available: ")
            return None
            
        except Exception as e:
            print(f"    ⚠️ Error on attempt {attempt + 1}: {e}")
            
            if attempt < max_retries - 1:
                # Exponential backoff for retries
                retry_delay = delay * (2 ** attempt) + random.uniform(1, 3)
                print(f"    Retrying in {retry_delay:.1f} seconds...")
                time.sleep(retry_delay)
            else:
                print(f"    ❌ Failed after {max_retries} attempts")
                return None
    
    return None

def save_to_docx(transcripts_dict, output_filename="data/merged_transcripts.docx"):
    """Save transcripts to Word document"""
    doc = Document()
    doc.add_heading("Merged Video Transcripts", 0)

    for url, transcript in transcripts_dict.items():
        doc.add_heading(f"Transcript for: {url}", level=1)
        if transcript:
            doc.add_paragraph(transcript)
        else:
            doc.add_paragraph("❌ No transcript available.")

    doc.save(output_filename)
    print(f"✅ Saved merged transcripts to: {output_filename}")

def save_log(bad_urls, log_filename="logs/transcript_log.txt"):
    """Save error log"""
    with open(log_filename, "w", encoding="utf-8") as f:
        for url, reason in bad_urls.items():
            f.write(f"{url} - ❌ {reason}\n")
    print(f"📝 Log saved to: {log_filename}")

def main(urls):
    """Main function using only Method 2 (Rate-limited API)"""


    transcripts = {}
    bad_urls = {}

    print("🚀 Starting transcript extraction using Method 2 (Rate-limited API)")
    print(f"📊 Processing {len(urls)} URLs with enhanced rate limiting...\n")

    for i, url in enumerate(urls):
        print(f"🔄 Processing {i+1}/{len(urls)}: {url}")
        
        video_id = extract_video_id(url)
        if not video_id:
            bad_urls[url] = "Invalid URL format or missing video ID"
            print(f"  ❌ Invalid URL format")
            continue

        print(f"  📹 Video ID: {video_id}")
        
        # Use Method 2 with enhanced rate limiting
        transcript = get_transcript_with_rate_limit(
            video_id, 
            delay=4,  # Increased base delay
            max_retries=3
        )
        
        if transcript:
            transcripts[url] = transcript
            print(f"  ✅ Success! Transcript length: {len(transcript)} characters\n")
        else:
            transcripts[url] = None
            bad_urls[url] = "Transcript not available or rate limited"
            print(f"  ❌ Failed to get transcript\n")
        
        # Add extra delay between videos to be safe
        if i < len(urls) - 1:  # Don't delay after the last URL
            print(f"  ⏳ Waiting 5 seconds before next video...")
            time.sleep(5)

    print("📄 Saving results...")
    save_to_docx(transcripts)
    
    if bad_urls:
        save_log(bad_urls)
        print(f"\n⚠️ {len(bad_urls)} URLs failed:")
        for url, reason in bad_urls.items():
            print(f"  - {url}: {reason}")
    
    successful_count = len([t for t in transcripts.values() if t is not None])
    print(f"\n🎉 Summary: {successful_count}/{len(urls)} transcripts extracted successfully!")

if __name__ == "__main__":
    main()