import time
import os
from docx import Document
from openai import OpenAI
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# === Set your OpenAI API key ===
openai_api_key = os.getenv("OPENAI_API_KEY")

client = OpenAI(api_key=openai_api_key)  # Replace with your actual key

# === GPT call using your updated function ===
def get_gpt_response(prompt, model="gpt-4.1-mini"):
    delay = 15
    print(f"⏳ Waiting {delay}s to respect rate limits...")
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        time.sleep(delay + 2)  # Respect TPM limit
        return response.choices[0].message.content
    except Exception as e:
        print("❌ GPT error:", e)
        return ""

# === Load .docx content ===
def load_docx_text(docx_path):
    doc = Document(docx_path)
    full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
    return "\n".join(full_text)

# === Save response to .docx ===
def save_to_docx(text, filename):
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)

# === Main process ===
def merge_video_summaries(docx_path):
    print(f"📂 Reading file: {docx_path}")
    doc_text = load_docx_text(docx_path)

    prompt = f"""
You are a professional tutor who has reviewed multiple video lecture summaries on the same topic. Your task is to:

1. Understand the key ideas across all summaries.
2. Identify repeated concepts, unique insights, and any contradictions or differences in explanations.
   - Clearly explain what the differences are and why they might exist.

Then merge all summaries into one well-structured explanation.

Organize the explanation like a tutor would:
- Start with Basics (definitions, foundations)
- Then Intermediate Concepts (examples, applications)
- Then Advanced Insights (expert-level analysis, edge cases)

Make the explanation easy to follow, detailed, and educational — like you're teaching a student who wants to fully understand the topic

Here are the video summaries:
{doc_text}

Now write the final, detailed, well-organized explanation """

    return prompt



def add_formatted_run(paragraph, text):
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*\*)')
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if pos < start:
            paragraph.add_run(text[pos:start])
        matched_text = match.group()
        cleaned_text = matched_text.strip('*')
        run = paragraph.add_run(cleaned_text)
        run.bold = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])



def write_to_word(text, output_file="generated_summary.docx"):
    doc = Document()
    title = doc.add_heading("Generated Topic-Wise Study Plan", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    lines = text.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif re.match(r"^Chapter \d+: ", line, re.IGNORECASE):
            doc.add_heading(line, level=2)
        elif line.startswith(("-", "*", "•")):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_run(p, line[1:].strip())
        elif re.match(r"^\d+\.", line):
            p = doc.add_paragraph(style="List Number")
            add_formatted_run(p, line.strip())
        else:
            p = doc.add_paragraph()
            add_formatted_run(p, line)

    doc.save(output_file)

if __name__ == "__main__":
    input_file = "merged_summary.docx"
    print("📤 Generating final prompt...")
    prompt = merge_video_summaries(input_file)

    print("🧠 Getting response from GPT...")
    merged_text = get_gpt_response(prompt)

    if merged_text:
        output_file_formatted = "generated_summary.docx"

        write_to_word(merged_text, output_file_formatted)

        print(f"✅ Formatted study plan saved to '{output_file_formatted}'")
    else:
        print("⚠️ No response generated from GPT.")


