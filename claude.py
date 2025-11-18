"""
Study Guide Converter - Convert Word documents to beautiful interactive HTML study guides
Author: AI Assistant
Requirements: pip install anthropic python-docx
"""
import os
import anthropic
import docx
import json
import sys
from pathlib import Path
from typing import Optional

class StudyGuideConverter:
    """Main converter class for transforming Word documents into study guides"""
    
    def __init__(self, api_key: str):
        """Initialize the converter with Anthropic API key"""
        self.client = anthropic.Anthropic(api_key=api_key)
        self.model = "claude-sonnet-4-20250514"
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text content from Word document
        
        Args:
            file_path: Path to the .docx file
            
        Returns:
            Extracted text as string
        """
        print(f"📖 Reading document: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            full_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    full_text.append(text)
            
            # Extract tables if any
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        full_text.append(row_text)
            
            content = "\n".join(full_text)
            print(f"✓ Extracted {len(content)} characters from document")
            return content
            
        except Exception as e:
            raise Exception(f"Error reading Word document: {str(e)}")
    
    # def generate_study_guide_html(self, document_text: str, course_name: Optional[str] = None) -> str:
    def generate_study_guide_html(
        self,
        document_text,
        course_name=None,
        theme="Light",
        font="Inter",
        detail_level="Medium",
        duration=4,
        tone="Professional"
    ):

        """
        Use Claude API to generate enhanced study guide with HTML
        
        Args:
            document_text: The extracted text from the document
            course_name: Optional course name for the title
            
        Returns:
            Complete HTML code as string
        """
        print("🤖 Generating enhanced study guide with Claude API...")
        
        course_instruction = f"\nThe course name is: {course_name}" if course_name else ""
        
        # prompt = f"""You are an expert educational content designer and web developer. Transform this study guide document into a stunning, modern, interactive HTML page.
        prompt = f"""
You are an expert educational content designer and professional web developer.

USER CUSTOMIZATION (apply strictly):
- Color Theme: {theme}
- Font Family: {font}
- Detail Level: {detail_level}
- Study Duration: {duration} weeks
- Writing Tone: {tone}

Apply these preferences in:
- HTML structure
- CSS styling
- Tone, phrasing, and level of detail
- Section length & density
- Color palette
- Typography
- Emphasis & layout spacing
- Box shadows, gradients, accents

Now transform the following document into a beautiful, interactive HTML study guide:

COURSE NAME (optional): {course_name}

DOCUMENT CONTENT:
{document_text}

IMPORTANT:
- All generated HTML must respect the user-selected theme, font, detail level, and tone.
- Tone = {tone} must influence writing style.
- Detail level = {detail_level} must control how verbose sections are.
- Theme = {theme} must influence colors.
- Font = {font} must be used throughout (import via Google Fonts).

{course_instruction}

Create a beautiful, comprehensive HTML study guide with these features:

1. **HERO SECTION**:
   - Eye-catching gradient background
   - Course title (large and bold)
   - Subtitle with key course info
   - Modern design with animations

2. **EXECUTIVE SUMMARY**:
   - Concise 3-4 sentence overview of the entire content
   - Highlight what students will learn
   - Use engaging language

3. **TABLE OF CONTENTS**:
   - Clickable navigation to all major sections
   - Smooth scroll behavior
   - Do NOT use sticky or fixed positioning
   - TOC should scroll normally with the page

4. **MAIN CONTENT SECTIONS**:
   - Organize content into logical chapters/topics
   - Each section with unique color accent
   - Use icons/emojis relevant to each topic
   - Include expandable/collapsible subsections
   - Add visual separators between sections

5. **KEY CONCEPTS CARDS**:
   - Highlight the most important concepts
   - Beautiful card design with hover effects
   - Color-coded by difficulty or topic

6. **INTERACTIVE ELEMENTS**:
   - Progress tracker with checkboxes for each major topic
   - "Mark as Complete" functionality
   - Progress percentage display
   - The progress bar must scroll normally with the page
   - Do NOT make the progress bar sticky or fixed
   - Tooltip hints on hover
   - Smooth transitions

7. **STUDY TIPS SECTION**:
   - AI-generated study recommendations
   - Time management tips (detailed)
   - Memory techniques relevant to the content

8. **QUICK REFERENCE**:
   - Glossary of key terms
   - Important formulas/definitions in cards
   - Quick facts section

9. **VISUAL DESIGN**:
   - Modern color scheme (use gradients: purple-blue, teal-green, or pink-orange)
   - Beautiful typography (multiple font weights)
   - Responsive design (mobile-friendly)
   - Smooth scroll behavior
   - Hover effects on interactive elements
   - Box shadows and depth
   - Professional spacing and layout

10. **EXTRAS**:
    - Print-friendly CSS
    - Notes section where students can add their own notes

TECHNICAL REQUIREMENTS:
- Return ONLY complete, valid HTML
- Include all CSS in <style> tags
- Include all JavaScript in <script> tags
- Make it completely self-contained (no external dependencies except fonts)
- Use Google Fonts (Poppins, Inter, or Roboto)
- Ensure all interactive features work
- Add smooth animations with CSS transitions
- Make it production-ready
- dont use any external libraries or frameworks
- use minimal javascript and limied dropdowns

Make this study guide so visually appealing and useful that students will actually want to study from it!"""

        try:
            message = self.client.messages.create(
                model=self.model,
                max_tokens=16000,
                temperature=1,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            html_content = message.content[0].text
            
            # Clean up the response (remove markdown code blocks if present)
            if html_content.startswith("```html"):
                html_content = html_content.split("```html")[1]
                html_content = html_content.rsplit("```")[0]
            elif html_content.startswith("```"):
                html_content = html_content.split("```")[1]
                html_content = html_content.rsplit("```")[0]
            
            print("✓ Study guide generated successfully")
            return html_content.strip()
            
        except anthropic.APIError as e:
            raise Exception(f"Claude API Error: {str(e)}")
    
    def save_html_file(self, html_content: str, output_path: str) -> None:
        """
        Save the generated HTML to a file
        
        Args:
            html_content: The HTML content to save
            output_path: Path where to save the file
        """
        print(f"💾 Saving HTML file to: {output_path}")
        
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            print(f"✅ Study guide saved successfully!")
        except Exception as e:
            raise Exception(f"Error saving HTML file: {str(e)}")
    
    # def convert(self, docx_path: str, output_path: str, course_name: Optional[str] = None) -> None:
    def convert(self, docx_path, output_path, course_name=None,
            theme="Light", font="Inter", 
            detail_level="Medium", duration=4, tone="Professional"):
        """
        Main conversion method
        
        Args:
            docx_path: Path to input Word document
            output_path: Path for output HTML file
            course_name: Optional course name
        """
        print("\n" + "="*60)
        print("📚 STUDY GUIDE CONVERTER")
        print("="*60 + "\n")
        
        # Step 1: Extract text
        document_text = self.extract_text_from_docx(docx_path)
        
        if not document_text:
            raise Exception("No text content found in the document")
        
        # Step 2: Generate HTML
        # html_content = self.generate_study_guide_html(document_text, course_name)
        html_content = self.generate_study_guide_html(
            document_text=document_text,
            course_name=course_name,
            theme=theme,
            font=font,
            detail_level=detail_level,
            duration=duration,
            tone=tone
        )

        
        # Step 3: Save file
        self.save_html_file(html_content, output_path)
        
        print("\n" + "="*60)
        print(f"🎉 DONE! Open '{output_path}' in your browser")
        print("="*60 + "\n")


def main(topic, theme, font, detail_level, duration, tone):
    """Main entry point"""

    API_KEY = os.getenv("ANTHROPIC_API_KEY")

    INPUT_FILE = "gpt_study_plan.docx"
    OUTPUT_FILE = "study_guide.html"
    COURSE_NAME = topic

    if not Path(INPUT_FILE).exists():
        print(f"❌ Error: Input file '{INPUT_FILE}' not found")
        sys.exit(1)

    try:
        converter = StudyGuideConverter(api_key=API_KEY)
        converter.convert(
            docx_path=INPUT_FILE,
            output_path=OUTPUT_FILE,
            course_name=COURSE_NAME,
            theme=theme,
            font=font,
            detail_level=detail_level,
            duration=duration,
            tone=tone
        )
    except Exception as e:
        print(f"\n❌ Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()