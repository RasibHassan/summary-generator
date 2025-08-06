from tavily import TavilyClient
from docx import Document
from datetime import datetime
import os
from dotenv import load_dotenv

load_dotenv()
TAVILY_API_KEY = os.getenv("TAVILY_API_KEY")
# Initialize Tavily client with your API key
tavily_client = TavilyClient(api_key=TAVILY_API_KEY)


def main(urls, answers=None):
    """Extract text from webpages and save to docx"""
    print("🌐 Fetching webpage content with Tavily...")

    response = tavily_client.extract(
        urls=urls,
        include_images=False,
        include_favicon=False,
        extract_depth="advanced",
        format="text"
    )

    # Create Word document
    doc = Document()
    doc.add_heading("Extracted Web Page Content", level=1)

    if answers and isinstance(answers, str):
        print("answers:", answers)
        doc.add_heading("Tavily Summary Answer", level=2)
        doc.add_paragraph(answers)
    
    for result in response.get("results", []):
        url = result.get("url", "Unknown URL")
        content = result.get("raw_content", "No content found.")
        doc.add_heading(url, level=2)
        clean_content = ''.join(c for c in content if c.isprintable() and c not in '\x00\x01\x02\x03\x04\x05\x06\x07\x08\x0b\x0c\x0e\x0f\x10\x11\x12\x13\x14\x15\x16\x17\x18\x19\x1a\x1b\x1c\x1d\x1e\x1f')
        doc.add_paragraph(clean_content)

    os.makedirs("data", exist_ok=True)
    doc_filename = "data/extracted_web_content.docx"
    doc.save(doc_filename)
    print(f"✅ All content saved to {doc_filename}")

    # Log failed URLs
    failed_urls = response.get("failed_results", [])
    if failed_urls:
        os.makedirs("logs", exist_ok=True)
        log_filename = "logs/failed_urls.log"
        with open(log_filename, "a", encoding="utf-8") as log_file:
            log_file.write(f"\n--- Log: {datetime.now()} ---\n")
            for failed in failed_urls:
                url = failed.get("url", "Unknown URL")
                reason = failed.get("reason", "No reason provided")
                log_file.write(f"{url} - Reason: {reason}\n")
        print(f"❌ {len(failed_urls)} failed URLs logged to {log_filename}")
    else:
        print("✅ All URLs processed successfully!")
