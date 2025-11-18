import os
import fitz  # PyMuPDF
import docx
from openai import OpenAI
from docx import Document
import sys
import re
import tiktoken
import time
from dotenv import load_dotenv

load_dotenv()
# === Save formatted Word document ===
def save_to_word(content, filename="gpt_study_plan.docx"):
    doc = Document()
    doc.add_heading("Study Plan", level=0)

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            continue
        elif line.startswith("### ") or line.startswith("## "):
            doc.add_heading(line.replace("#", "").strip(), level=1)
        elif line.startswith("- "):
            doc.add_paragraph(line, style='List Bullet')
        elif re.match(r"^\d+\.", line):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(filename)


# === OpenAI Client Setup ===
openai_api_key = os.getenv("OPENAI_API_KEY")
client = OpenAI(api_key=openai_api_key)  # Replace with your actual key


# === File Extraction ===
def extract_text_from_pdf(file_path):
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        return ""


def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        return ""


def load_all_text(folder_path):
    all_text = ""
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if filename.endswith(".pdf"):
            all_text += extract_text_from_pdf(file_path) + "\n"
        elif filename.endswith(".docx"):
            all_text += extract_text_from_docx(file_path) + "\n"
    return all_text.strip()


# === Clean and Sanitize Text ===
def clean_text(text):
    text = ''.join(c for c in text if c.isprintable() and c not in '\x00')
    return " ".join(text.split())


# === Count Tokens Using tiktoken ===
def count_tokens(text):
    enc = tiktoken.encoding_for_model('gpt-4')
    return len(enc.encode(text))


# === Prompt Builder ===
def build_prompt(certificate_name, context):
    return f"""

You are an expert educational mentor specializing in creating comprehensive, actionable study guides for **{certificate_name}**.

## Your Input
You will receive unstructured content from multiple sources:
- YouTube video transcripts (may be auto-generated with errors)
- Reddit posts and discussions
- Blog posts and articles
- Forum discussions (Quora, Stack Exchange, etc.)
- User experiences and testimonials

**Note:** Content may be fragmented, repetitive, or contain contradictions. Your job is to synthesize, organize, and extract the most valuable insights.

---

## Your Mission
Transform raw, unstructured content into a **detailed, exam-focused study guide** that helps users understand:
1. **WHAT** to study (key topics and concepts)
2. **WHY** it matters (exam relevance and real-world application)
3. **HOW** to study it effectively (strategies, resources, examples)

**DO NOT** create a generic overview. Extract specific, actionable content that directly helps with exam preparation or subject mastery.

---

## Output Structure

### 1. 🎯 Most Important Topics to Focus On
- Identify the **3-7 highest-impact topics** mentioned across sources
- For each topic, provide:
  - **Why it's critical:** Explain its weight in the exam or importance in the field
  - **What makes it challenging:** Common difficulties or misconceptions
  - **Frequency in exam:** How often it appears (if mentioned in context)

*Use short paragraphs, not bullet lists. Make it conversational but information-dense.*

---

### 2. 📚 Core Concepts & Learning Content

Break down each major topic with:

#### Topic Name
**Concept Explanation:** Clear, detailed explanation of the concept. Include definitions, mechanisms, and how it works.

**Practical Examples:** Real-world scenarios or exam-style examples that illustrate the concept. Use analogies where helpful.

**Common Formulas/Rules:** Any mathematical formulas, mnemonics, or rules of thumb. Explain when and how to apply them.

**Tricky Points:** Specific aspects that confuse learners. Clarify misconceptions and provide memory aids.

**Connection to Other Topics:** How this concept links to other areas (helps with holistic understanding).

*Repeat this structure for all major topics extracted from the content.*

---

### 3. 💡 Why This Matters: Exam Relevance

For each concept covered, explain:
- **Exam question types:** How this topic typically appears (multiple choice, case study, essay, practical)
- **Scoring impact:** Is this a high-point topic? Does it appear in multiple sections?
- **Real-world application:** How professionals use this knowledge (motivates learning)
- **Red flags:** What incorrect answers or approaches to avoid

*This section helps students prioritize and understand the "so what?" of each topic.*

---

### 4. 📖 Curated Resources

List the most valuable resources mentioned in the context:

**Books:**
- Title, Author - Why it's recommended, what chapters/sections to focus on

**Online Courses/Videos:**
- Creator/Platform, Link (if available) - What makes it effective, specific topics covered

**Practice Materials:**
- Mock exams, question banks, flashcard sets - Where to find them, how to use them

**Tools & Software:**
- Any apps, calculators, simulation tools - Specific use cases

**Communities:**
- Forums, Discord servers, study groups - How they help

*Only include resources explicitly mentioned or strongly implied in the context. Don't invent recommendations.*

---

### 5. 🧠 Effective Study Strategies

Extract and organize preparation techniques shared by successful candidates:

**Time Management:**
- How long to study (total hours, weeks before exam)
- Daily/weekly schedule suggestions
- When to start practicing vs learning concepts

**Study Techniques:**
- Active recall methods mentioned
- Spaced repetition strategies
- Note-taking approaches
- Group study recommendations

**Practice Strategy:**
- How many practice exams to take
- When to start practice tests
- How to review incorrect answers
- Simulation vs real exam differences

**Week-by-Week Approach:**
- Phase 1 (Weeks 1-X): Focus areas
- Phase 2 (Weeks X-Y): Focus areas
- Final Week: Focus areas

*Base this entirely on actual user experiences and recommendations from the content.*

---

### 6. ⚠️ User Insights: What Works & What Doesn't

**What Successful Candidates Did:**
- Specific habits and routines that led to success
- Mindset approaches that helped
- Resources they found most valuable

**Common Mistakes to Avoid:**
- What tripped people up
- Topics that are over/under-studied
- Timing issues (too much/little prep)
- Resource quality issues

**Exam Day Reality:**
- What the actual exam was like vs expectations
- Time pressure insights
- Question difficulty compared to practice materials
- Surprising elements people encountered

**Emotional/Mental Preparation:**
- Dealing with anxiety
- Confidence building
- What helped them stay motivated

*This section should feel like advice from someone who just passed the exam.*

---

### 7. 📅 Suggested Study Plan

Based on all extracted information, provide a **realistic, detailed study schedule**:

**Assumptions:**
- Study time available per week: [X hours based on context, or suggest range]
- Total preparation time: [Y weeks/months]
- Prior knowledge level: [Beginner/Intermediate/Advanced - if mentioned]

**Weekly Breakdown:**

**Week 1-2: Foundation Building**
- Topics to cover: [Specific topics]
- Resources to use: [Specific materials]
- Goals: [Measurable objectives]
- Practice: [What type, how much]

**Week 3-4: Deep Dive**
[Continue pattern]

**[Continue for full preparation timeline]**

**Final Week: Exam Readiness**
- Review strategy
- Practice exam schedule
- What to avoid (no new topics)
- Day-before preparation

---

## Important Guidelines

**Quality Standards:**
- ✅ Extract specific, actionable information
- ✅ Use clear, engaging language
- ✅ Provide context and reasoning, not just facts
- ✅ Include examples and practical applications
- ❌ Don't invent information not in the context
- ❌ Don't use excessive bullet points (prefer paragraphs)
- ❌ Don't be generic or vague
- ❌ Don't include sections with no relevant data

**Handling Missing Information:**
- If a section has no relevant data from the context, **skip that section entirely**
- Never say "no information provided" - just omit the section
- Focus depth on areas where you have rich content

**Tone:**
- Knowledgeable but approachable
- Like a mentor who's been through this
- Encouraging but realistic
- Specific and detail-oriented

---

## Context to Analyze

///
{context}
///
---

**Now generate the comprehensive study guide based on the context provided above.**

"""


# === GPT Request ===
def split_text_by_tokens(text, max_tokens=280000, model="gpt-4.1-mini"):
    enc = tiktoken.encoding_for_model('gpt-4')
    tokens = enc.encode(text)
    chunks = []
    for i in range(0, len(tokens), max_tokens):
        chunk = enc.decode(tokens[i:i + max_tokens])
        chunks.append(chunk)
    return chunks


def get_gpt_response(prompt, model="gpt-4.1-mini", max_tpm=400000):
    token_count = count_tokens(prompt)
    delay = 15
    print(f"⏳ Sleeping for {delay:.2f} seconds to respect TPM rate limit")

    try:
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
        )
        time.sleep(delay + 2)  # Add buffer
        return response.choices[0].message.content
    except Exception as e:
        print("❌ Error:", str(e).encode('utf-8', errors='ignore').decode())

        return ""


def main(certificate_name="Certificate"):
    folder_path = "data"
    raw_text = load_all_text(folder_path)
    if not raw_text:
        return

    context = clean_text(raw_text)
    print(f"Processing {certificate_name} with context length: {len(context)} characters")
    total_tokens = count_tokens(context)
    print(f"🔢 Total tokens: {total_tokens}")

    chunks = split_text_by_tokens(context, max_tokens=280000)
    print(f"✂️ Total chunks: {len(chunks)} (each ≤ 950k tokens)")

    final_output = ""

    for i, chunk in enumerate(chunks):
        print(f"\n🚀 Sending chunk {i+1}/{len(chunks)} to GPT...")

        prompt = build_prompt(certificate_name, chunk)
        result = get_gpt_response(prompt)
        if not result:
            print("⚠️ Empty response from GPT.")
        final_output += result + "\n\n"
    print(f"Total tokens processed: {total_tokens}")
    print(f"Final output length: {len(final_output)} characters")
    save_to_word(final_output.strip())


if __name__ == "__main__":
    main()
